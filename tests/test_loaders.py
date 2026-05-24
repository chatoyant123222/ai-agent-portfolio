import unittest
import os
import sys
import tempfile
import shutil
from pathlib import Path

# 将 src 加入路径
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from docx import Document as DocxDocument
from loaders.pdf_loader import PDFLoader
from loaders.word_loader import WordLoader
from models.document import Document, DocumentLoader


class TestDocumentLoaderBase(unittest.TestCase):
    """测试抽象基类约束"""

    def test_cannot_instantiate_abc(self):
        with self.assertRaises(TypeError):
            DocumentLoader()


class TestPDFLoader(unittest.TestCase):
    def setUp(self):
        self.loader = PDFLoader()

    def test_file_not_found(self):
        with self.assertRaises(FileNotFoundError):
            self.loader.load("不存在的文件.pdf")

    def test_invalid_path(self):
        with self.assertRaises(FileNotFoundError):
            self.loader.load("")


class TestWordLoader(unittest.TestCase):
    def setUp(self):
        self.loader = WordLoader()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _create_docx(self, filename: str, paragraphs: list, title: str = None, author: str = None):
        path = os.path.join(self.temp_dir, filename)
        doc = DocxDocument()
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    def test_load_valid_docx(self):
        path = self._create_docx(
            "test.docx",
            ["这是第一段", "这是第二段", "RAG 是一种检索增强生成技术。"],
            title="测试文档",
            author="单元测试"
        )

        doc = self.loader.load(path)

        self.assertIsInstance(doc, Document)
        self.assertEqual(doc.source, path)
        self.assertIn("这是第一段", doc.content)
        self.assertIn("RAG 是一种检索增强生成技术。", doc.content)
        self.assertEqual(doc.metadata["author"], "单元测试")
        self.assertEqual(doc.metadata["title"], "测试文档")
        self.assertEqual(doc.metadata["paragraph_count"], 3)

    def test_load_with_table(self):
        path = os.path.join(self.temp_dir, "table.docx")
        docx = DocxDocument()
        docx.add_paragraph("文档开头")
        table = docx.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "姓名"
        table.rows[0].cells[1].text = "年龄"
        table.rows[1].cells[0].text = "张三"
        table.rows[1].cells[1].text = "25"
        docx.save(path)

        doc = self.loader.load(path)
        self.assertIn("文档开头", doc.content)
        self.assertIn("表格内容", doc.content)
        self.assertEqual(doc.metadata["table_count"], 1)

    def test_file_not_found(self):
        with self.assertRaises(FileNotFoundError):
            self.loader.load("不存在的文件.docx")

    def test_invalid_extension(self):
        fake = os.path.join(self.temp_dir, "fake.txt")
        with open(fake, "w", encoding="utf-8") as f:
            f.write("假文件")
        with self.assertRaises(ValueError):
            self.loader.load(fake)


if __name__ == "__main__":
    unittest.main(verbosity=2)
