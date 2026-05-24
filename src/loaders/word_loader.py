import os
import uuid
import sys
from pathlib import Path
from typing import Any, Optional
from docx import Document as DocxDocument

sys.path.append(str(Path(__file__).parent.parent))
from models.document import Document, DocumentLoader


class WordLoader(DocumentLoader):
    """Word 文档加载器，支持 .docx 格式"""

    def __init__(self, extract_metadata: bool = True, extract_tables: bool = True):
        self.extract_metadata = extract_metadata
        self.extract_tables = extract_tables

    def load(self, source: str) -> Document:
        if not os.path.exists(source):
            raise FileNotFoundError(f"Word 文件不存在: {source}")
        if not source.lower().endswith('.docx'):
            raise ValueError(f"不支持的文件格式: {source}，仅支持 .docx")

        try:
            docx = DocxDocument(source)
            return self.parse(docx, source_path=source)
        except Exception as e:
            raise ValueError(f"读取 Word 文件失败: {e}")

    def parse(self, raw_data: Any, source_path: Optional[str] = None) -> Document:
        if isinstance(raw_data, str):
            docx = DocxDocument(raw_data)
            source_path = raw_data
        else:
            docx = raw_data

        # 提取段落文本
        paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
        full_content = "\n".join(paragraphs)

        # 提取表格内容（可选）
        if self.extract_tables and docx.tables:
            tables_text = []
            for table in docx.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables_text.append("\n".join(rows))
            full_content += "\n\n[表格内容]\n" + "\n\n".join(tables_text)

        # 构建元数据
        metadata = {
            "paragraph_count": len(paragraphs),
            "table_count": len(docx.tables),
        }

        if self.extract_metadata:
            core = docx.core_properties
            metadata.update({
                "title": core.title,
                "author": core.author,
                "subject": core.subject,
                "created": str(core.created) if core.created else None,
                "modified": str(core.modified) if core.modified else None,
            })

        return Document(
            id=str(uuid.uuid4()),
            content=full_content,
            metadata=metadata,
            source=source_path or "unknown"
        )
